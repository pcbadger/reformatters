import os
import sys
from docx import Document ## python-docx
from docx.shared import Mm
from docx.shared import Cm

from PIL import Image ## pillow

directory_in_str = "pics"
directory = os.fsencode(directory_in_str)
fileType = ".jpg"

document = Document()

if sys.argv[1]:
    outputDoc = sys.argv[1] + ".docx"
else:
    outputDoc = 'docx_file.docx'

try:
    os.remove(outputDoc)
except:
    print("Nothing to remove")

# def getDocWidth(document):
#     section = document.sections[0]
#     return (section.page_width - section.left_margin - section.right_margin) / 36000

def convertImage(image):
    convertedImage = Image.open(image)
    convertedImage.save("tmp.png")

def addImage(image, outputDoc):
    #document.add_picture(image, Mm(getDocWidth(outputDoc)))
    document.add_picture(image, height=Cm(20.5), width=Cm(14.5))
    print("Appending " + image + " to " + outputDoc)
    document.add_page_break()

for file in os.listdir(directory):
    fileName = os.fsdecode(file)
    if fileName.endswith(fileType):
        filePath = directory_in_str + "/" + fileName
        with open(filePath) as infile:

            try:
                addImage(filePath, outputDoc)
            except:
                print("Unable to add  " + filePath )
                convertImage(filePath)
                addImage("tmp.png", outputDoc)
                os.remove("tmp.png")
        continue
    else:
        continue
document.save(outputDoc)
