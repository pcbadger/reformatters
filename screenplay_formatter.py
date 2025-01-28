from docx import Document
from docx.shared import Pt, Inches

def format_screenplay(input_file, word_output_file, text_output_file):
    # Create a new Word document
    doc = Document()

    # Define styles for screenplay elements
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Courier New'
    normal_style.font.size = Pt(12)

    name_style = doc.styles.add_style('NameStyle', 1)
    name_style.font.name = 'Courier New'
    name_style.font.size = Pt(12)
    name_style.font.bold = True

    dialog_style = doc.styles.add_style('DialogStyle', 1)
    dialog_style.font.name = 'Courier New'
    dialog_style.font.size = Pt(12)

    # Read input text from file
    with open(input_file, 'r') as file:
        lines = file.readlines()

    # Clean up lines
    cleaned_lines = []
    blank_line_count = 0

    for line in lines:
        stripped_line = line.rstrip()  # Remove trailing whitespace

        if stripped_line == "":  # Check for blank line
            blank_line_count += 1
        else:
            blank_line_count = 0

        if blank_line_count <= 1:  # Allow at most one blank line
            cleaned_lines.append(stripped_line)

    # Save cleaned text to a plain text file
    with open(text_output_file, 'w') as text_file:
        text_file.write("\n".join(cleaned_lines) + "\n")

    prev_line_type = None  # Tracks the type of the previous line ("NAME", "DIALOG", or "ACTION")

    for line in cleaned_lines:
        line = line.strip()

        if not line:  # Skip completely blank lines
            continue

        # Check if line is a NAME (single word not following another NAME)
        if len(line.split()) == 1 and prev_line_type != "NAME":
            if prev_line_type == "ACTION":
                doc.add_paragraph(" ", style=normal_style)  # Blank line before NAME

            paragraph = doc.add_paragraph(line.upper(), style='NameStyle')
            paragraph.alignment = 1  # Center alignment
            prev_line_type = "NAME"

        # Check if line is DIALOG (follows a NAME)
        elif prev_line_type == "NAME":
            paragraph = doc.add_paragraph(line, style='DialogStyle')
            paragraph.alignment = 3  # Justified
            paragraph.paragraph_format.left_indent = Inches(1)  # 1-inch margins
            paragraph.paragraph_format.right_indent = Inches(1)
            doc.add_paragraph(" ", style=normal_style)  # Blank line after DIALOG
            prev_line_type = "DIALOG"

        # Otherwise, it's ACTION
        else:
            if prev_line_type in ["NAME", "DIALOG"]:
                doc.add_paragraph(" ", style=normal_style)  # Blank line before ACTION

            paragraph = doc.add_paragraph(line, style='Normal')
            doc.add_paragraph(" ", style=normal_style)  # Blank line after ACTION
            prev_line_type = "ACTION"

    # Remove paragraphs with only a space
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "":
            p = paragraph._element
            p.getparent().remove(p)

    # Save the formatted screenplay to a Word document
    doc.save(word_output_file)

# Example usage
input_file = "screenplay_input.txt"  # Input text file
word_output_file = "screenplay.docx"  # Word output file
text_output_file = "screenplay_cleaned.txt"  # Plain text output file

format_screenplay(input_file, word_output_file, text_output_file)
print(f"Screenplay saved to {word_output_file} and {text_output_file}")